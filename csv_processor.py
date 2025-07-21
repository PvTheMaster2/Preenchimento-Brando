#!/usr/bin/env python3
"""
CSV Processor - Sistema de Processamento de Planilha Preenchida
Processa arquivo CSV com informações dos shifts e gera documentos Word profissionais
"""

import pandas as pd
import json
import random
from datetime import datetime, timedelta
from docxtpl import DocxTemplate
import argparse
import os
from pathlib import Path

class CSVProcessor:
    def __init__(self, config_path="config.json", content_pools_path="content_pools.json"):
        """
        Inicializa o processador CSV
        
        Args:
            config_path: Caminho para arquivo de configuração
            content_pools_path: Caminho para pools de conteúdo automático
        """
        self.config = self.load_config(config_path)
        self.content_pools = self.load_content_pools(content_pools_path)
        
    def load_config(self, config_path):
        """Carrega configurações do projeto"""
        try:
            with open(config_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except FileNotFoundError:
            # Configuração padrão se arquivo não existir
            return {
                "student_name": "Student Name",
                "supervisor": "Supervisor Name",
                "course": "SIT40521 Certificate IV in Kitchen Management",
                "template_path": "template_logbook.docx"
            }
    
    def load_content_pools(self, content_pools_path):
        """Carrega pools de conteúdo para preenchimento automático"""
        try:
            with open(content_pools_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except FileNotFoundError:
            # Pools básicos se arquivo não existir
            return {
                "food_details": [
                    "Preparei hambúrgueres, batatas fritas e saladas do dia.",
                    "Trabalhei com grelhados, risotos e pratos principais.",
                    "Foquei em preparação de sopas, massas e sobremesas."
                ],
                "special_requests": [
                    "Cliente pediu refeição sem glúten.",
                    "Solicitação especial de temperatura na carne.",
                    "Adaptação para cliente vegetariano."
                ],
                "problems": [
                    "Equipamento apresentou problema temporário.",
                    "Demanda maior que esperado durante pico.",
                    "Ajuste necessário na temperatura de cozimento."
                ],
                "solutions": [
                    "Resolvi rapidamente com ajuda da equipe.",
                    "Ajustei processo e melhorei eficiência.",
                    "Comuniquei problema e implementei solução."
                ],
                "learnings": [
                    "Aprendi importância do trabalho em equipe.",
                    "Desenvolvi técnicas de organização no tempo.",
                    "Melhorei comunicação com colegas."
                ]
            }
    
    def load_csv_data(self, csv_path):
        """
        Carrega dados do CSV preenchido
        
        Args:
            csv_path: Caminho para arquivo CSV
            
        Returns:
            DataFrame com dados dos shifts
        """
        try:
            df = pd.read_csv(csv_path, encoding='utf-8')
            print(f"✅ CSV carregado: {len(df)} shifts encontrados")
            return df
        except Exception as e:
            print(f"❌ Erro ao carregar CSV: {e}")
            return None
    
    def fill_missing_data(self, df):
        """
        Preenche campos vazios com conteúdo dos pools
        
        Args:
            df: DataFrame com dados dos shifts
            
        Returns:
            DataFrame com todos os campos preenchidos
        """
        filled_df = df.copy()
        
        for index, row in filled_df.iterrows():
            # Preenche food_details se vazio
            if pd.isna(row['food_details']) or row['food_details'] == '':
                filled_df.at[index, 'food_details'] = random.choice(self.content_pools['food_details'])
            
            # Preenche special_requests se vazio
            if pd.isna(row['special_requests']) or row['special_requests'] == '':
                filled_df.at[index, 'special_requests'] = random.choice(self.content_pools['special_requests'])
            
            # Preenche complaints_problems se vazio
            if pd.isna(row['complaints_problems']) or row['complaints_problems'] == '':
                filled_df.at[index, 'complaints_problems'] = random.choice(self.content_pools['problems'])
            
            # Preenche solutions_implemented se vazio
            if pd.isna(row['solutions_implemented']) or row['solutions_implemented'] == '':
                filled_df.at[index, 'solutions_implemented'] = random.choice(self.content_pools['solutions'])
            
            # Preenche debrief_learnings se vazio
            if pd.isna(row['debrief_learnings']) or row['debrief_learnings'] == '':
                filled_df.at[index, 'debrief_learnings'] = random.choice(self.content_pools['learnings'])
        
        return filled_df
    
    def generate_documents(self, df, output_folder="output_documents"):
        """
        Gera documentos Word para cada shift
        
        Args:
            df: DataFrame com dados completos dos shifts
            output_folder: Pasta de saída para documentos
        """
        # Cria pasta de saída se não existir
        Path(output_folder).mkdir(exist_ok=True)
        
        successful_docs = 0
        
        for index, row in df.iterrows():
            try:
                # Contexto para o template
                context = {
                    'student_name': self.config['student_name'],
                    'supervisor': self.config['supervisor'],
                    'course': self.config['course'],
                    'shift_number': row['shift_number'],
                    'date': row['date'],
                    'shift_type': row['shift_type'].title(),
                    'start_time': row['start_time'],
                    'end_time': row['end_time'],
                    'food_details': row['food_details'],
                    'special_requests': row['special_requests'],
                    'complaints_problems': row['complaints_problems'],
                    'solutions_implemented': row['solutions_implemented'],
                    'debrief_learnings': row['debrief_learnings']
                }
                
                # Nome do arquivo de saída
                output_filename = f"shift_{row['shift_number']:02d}_{row['date']}_{row['shift_type']}.docx"
                output_path = os.path.join(output_folder, output_filename)
                
                # Gera documento (usando template simples se não existir template específico)
                self.create_document(context, output_path)
                successful_docs += 1
                
            except Exception as e:
                print(f"❌ Erro ao gerar documento para shift {row['shift_number']}: {e}")
        
        print(f"✅ {successful_docs} documentos gerados com sucesso em '{output_folder}'")
    
    def create_document(self, context, output_path):
        """
        Cria documento Word individual
        
        Args:
            context: Dados do shift
            output_path: Caminho de saída do documento
        """
        # Template básico se não existir arquivo específico
        if not os.path.exists(self.config.get('template_path', 'template_logbook.docx')):
            self.create_basic_document(context, output_path)
        else:
            # Usa template existente
            doc = DocxTemplate(self.config['template_path'])
            doc.render(context)
            doc.save(output_path)
    
    def create_basic_document(self, context, output_path):
        """
        Cria documento básico sem template
        
        Args:
            context: Dados do shift
            output_path: Caminho de saída
        """
        from docx import Document
        
        doc = Document()
        
        # Título
        title = doc.add_heading(f'Logbook Entry - Shift {context["shift_number"]}', 0)
        
        # Informações básicas
        doc.add_paragraph(f"Student: {context['student_name']}")
        doc.add_paragraph(f"Course: {context['course']}")
        doc.add_paragraph(f"Supervisor: {context['supervisor']}")
        doc.add_paragraph(f"Date: {context['date']}")
        doc.add_paragraph(f"Shift Type: {context['shift_type']}")
        doc.add_paragraph(f"Time: {context['start_time']} - {context['end_time']}")
        
        # Seções principais
        sections = [
            ('Food Details', context['food_details']),
            ('Special Requests', context['special_requests']),
            ('Problems/Complaints', context['complaints_problems']),
            ('Solutions Implemented', context['solutions_implemented']),
            ('Debrief & Learnings', context['debrief_learnings'])
        ]
        
        for section_title, content in sections:
            doc.add_heading(section_title, level=1)
            doc.add_paragraph(content)
        
        doc.save(output_path)
    
    def generate_summary_report(self, df, output_path="summary_report.txt"):
        """
        Gera relatório resumo do processamento
        
        Args:
            df: DataFrame processado
            output_path: Caminho para arquivo de relatório
        """
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("RELATÓRIO DE PROCESSAMENTO CSV\n")
            f.write("=" * 50 + "\n\n")
            
            f.write(f"Total de shifts processados: {len(df)}\n")
            f.write(f"Data de processamento: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
            
            # Contagem por tipo de shift
            lunch_count = len(df[df['shift_type'] == 'lunch'])
            dinner_count = len(df[df['shift_type'] == 'dinner'])
            f.write(f"Shifts de almoço: {lunch_count}\n")
            f.write(f"Shifts de jantar: {dinner_count}\n\n")
            
            # Estatísticas de preenchimento manual vs automático
            manual_count = 0
            auto_count = 0
            
            for _, row in df.iterrows():
                if any(row[col] in self.content_pools.get(key, []) for col, key in [
                    ('food_details', 'food_details'),
                    ('special_requests', 'special_requests'),
                    ('complaints_problems', 'problems'),
                    ('solutions_implemented', 'solutions'),
                    ('debrief_learnings', 'learnings')
                ]):
                    auto_count += 1
                else:
                    manual_count += 1
            
            f.write(f"Shifts preenchidos manualmente: {manual_count}\n")
            f.write(f"Shifts completados automaticamente: {auto_count}\n")
            
        print(f"📊 Relatório gerado: {output_path}")

def main():
    """Função principal do script"""
    parser = argparse.ArgumentParser(description='Processar CSV de shifts e gerar documentos')
    parser.add_argument('--csv', required=True, help='Caminho para arquivo CSV preenchido')
    parser.add_argument('--output', default='output_documents', help='Pasta de saída para documentos')
    parser.add_argument('--config', default='config.json', help='Arquivo de configuração')
    parser.add_argument('--pools', default='content_pools.json', help='Arquivo de pools de conteúdo')
    
    args = parser.parse_args()
    
    print("🚀 Iniciando processamento de CSV...")
    
    # Inicializa processador
    processor = CSVProcessor(args.config, args.pools)
    
    # Carrega dados CSV
    df = processor.load_csv_data(args.csv)
    if df is None:
        return
    
    # Preenche campos vazios
    print("🔄 Preenchendo campos vazios automaticamente...")
    df_filled = processor.fill_missing_data(df)
    
    # Gera documentos
    print("📄 Gerando documentos Word...")
    processor.generate_documents(df_filled, args.output)
    
    # Gera relatório
    print("📊 Gerando relatório resumo...")
    processor.generate_summary_report(df_filled)
    
    print("✅ Processamento concluído com sucesso!")
    print(f"📁 Documentos salvos em: {args.output}")

if __name__ == "__main__":
    main() 
