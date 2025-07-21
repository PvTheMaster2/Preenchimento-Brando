#!/usr/bin/env python3
"""
CSV Processor Completo - Sistema de Processamento de Planilha Completa
Processa arquivo CSV completo com workflow e todos os campos do logbook
"""

import pandas as pd
import json
import random
from datetime import datetime, timedelta
from docxtpl import DocxTemplate
import argparse
import os
from pathlib import Path

class CSVProcessorCompleto:
    def __init__(self, config_path="config.json", content_pools_path="content_pools.json"):
        """
        Inicializa o processador CSV completo
        
        Args:
            config_path: Caminho para arquivo de configuração
            content_pools_path: Caminho para pools de conteúdo automático
        """
        self.config = self.load_config(config_path)
        self.content_pools = self.load_content_pools(content_pools_path)
        
    def load_config(self, config_path):
        """Carrega configurações do projeto"""
        try:
            with open(config_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except FileNotFoundError:
            return {
                "student_name": "Student Name",
                "supervisor": "Supervisor Name",
                "course": "SIT40521 Certificate IV in Kitchen Management",
                "establishment": "Fatcow on James Street",
                "menu_style": "à la carte",
                "template_path": "template_logbook_completo.docx"
            }
    
    def load_content_pools(self, content_pools_path):
        """Carrega pools de conteúdo para preenchimento automático"""
        try:
            with open(content_pools_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except FileNotFoundError:
            return {
                "prepared_for_service": [
                    "Preparei estação verificando equipamentos, organizei ingredientes e limpei superfícies.",
                    "Organizei mise en place, testei temperatura dos fornos e preparei utensílios.",
                    "Verifiquei inventário, preparei ingredientes frescos e organizei estação de trabalho."
                ],
                "special_requests": [
                    "Cliente pediu refeição sem glúten e molho à parte.",
                    "Solicitação especial de temperatura na carne e vegetais extras.",
                    "Adaptação para cliente vegetariano com substituição de proteína."
                ],
                "food_details": [
                    "Preparei hambúrgueres artesanais, batatas fritas e saladas frescas.",
                    "Trabalhei com grelhados, risotos e pratos principais da temporada.",
                    "Foquei em preparação de sopas, massas e sobremesas especiais."
                ],
                "complaints_problems": [
                    "Cliente reclamou sobre temperatura do prato.",
                    "Demanda maior que esperado durante horário de pico.",
                    "Equipamento apresentou problema temporário."
                ],
                "solutions_implemented": [
                    "Refiz o prato imediatamente e ajustei procedimentos.",
                    "Organizei melhor o fluxo de trabalho e comuniquei com a equipe.",
                    "Resolveu rapidamente com ajuda da equipe técnica."
                ],
                "debrief_learnings": [
                    "Aprendi importância do controle de temperatura e trabalho em equipe.",
                    "Desenvolvi técnicas de organização e comunicação eficaz.",
                    "Melhorei comunicação com garçons e coordenação de pratos."
                ],
                "handover_completed": [
                    "Sim, passei todas as informações para o turno seguinte.",
                    "Concluído com sucesso, informei sobre pedidos especiais pendentes.",
                    "Handover realizado, reportei status dos equipamentos e estoque."
                ],
                "customer_feedback": [
                    "Clientes elogiaram qualidade dos ingredientes e apresentação.",
                    "Feedback positivo sobre rapidez no atendimento e sabor dos pratos.",
                    "Comentários positivos sobre resolução rápida de problemas."
                ],
                "workflow_timeline": ["10:30", "11:00", "12:00", "13:00", "14:00", "15:00", "15:30", "16:00"],
                "workflow_tasks": [
                    "Check equipment and station setup",
                    "Mise en place preparation",
                    "Service preparation - lunch rush",
                    "Peak service management",
                    "Order completion and quality check",
                    "Station cleaning and sanitization",
                    "Handover preparation",
                    "Final checks and departure"
                ],
                "workflow_equipment": [
                    "Grill / Fryer / Prep surfaces",
                    "Knives / Cutting boards / Containers",
                    "All cooking equipment / Safety gloves",
                    "Thermometer / Plating equipment",
                    "Cleaning supplies / Sanitizer",
                    "Clipboards / Logbooks",
                    "All equipment shutdown"
                ],
                "workflow_communication": [
                    "Head Chef about daily specials",
                    "Sous chef about ingredient availability",
                    "Wait staff about special requirements",
                    "Kitchen team about order priorities",
                    "Head chef about food quality",
                    "Team about equipment status",
                    "Next shift about ongoing orders",
                    "Supervisor about shift completion"
                ]
            }
    
    def load_csv_data(self, csv_path):
        """Carrega dados do CSV preenchido"""
        try:
            df = pd.read_csv(csv_path, encoding='utf-8')
            print(f"✅ CSV carregado: {len(df)} shifts encontrados")
            return df
        except Exception as e:
            print(f"❌ Erro ao carregar CSV: {e}")
            return None
    
    def fill_missing_data(self, df):
        """Preenche campos vazios com conteúdo dos pools"""
        filled_df = df.copy()
        
        for index, row in filled_df.iterrows():
            # Preencher campos básicos
            if pd.isna(row['establishment']) or row['establishment'] == '':
                filled_df.at[index, 'establishment'] = self.config.get('establishment', 'Restaurant Name')
            
            if pd.isna(row['menu_style']) or row['menu_style'] == '':
                filled_df.at[index, 'menu_style'] = self.config.get('menu_style', 'à la carte')
            
            # Preencher campos de conteúdo
            content_fields = {
                'prepared_for_service': 'prepared_for_service',
                'special_requests': 'special_requests',
                'food_details': 'food_details',
                'complaints_problems': 'complaints_problems',
                'solutions_implemented': 'solutions_implemented',
                'debrief_learnings': 'debrief_learnings',
                'handover_completed': 'handover_completed',
                'customer_feedback': 'customer_feedback'
            }
            
            for field, pool_key in content_fields.items():
                if pd.isna(row[field]) or row[field] == '':
                    filled_df.at[index, field] = random.choice(self.content_pools[pool_key])
            
            # Preencher workflow se vazio
            for i in range(1, 9):  # 8 linhas de workflow
                timeline_col = f'workflow_timeline_{i}'
                task_col = f'workflow_task_{i}'
                equipment_col = f'workflow_equipment_{i}'
                communication_col = f'workflow_communication_{i}'
                
                if pd.isna(row[timeline_col]) or row[timeline_col] == '':
                    # Ajustar horários baseado no shift type
                    if row['shift_type'] == 'lunch':
                        base_times = ["10:30", "11:00", "12:00", "13:00", "14:00", "15:00", "15:30", "16:00"]
                    else:  # dinner
                        base_times = ["16:00", "17:00", "18:00", "19:00", "20:00", "21:00", "21:30", "22:00"]
                    
                    if i <= len(base_times):
                        filled_df.at[index, timeline_col] = base_times[i-1]
                        filled_df.at[index, task_col] = random.choice(self.content_pools['workflow_tasks'])
                        filled_df.at[index, equipment_col] = random.choice(self.content_pools['workflow_equipment'])
                        filled_df.at[index, communication_col] = random.choice(self.content_pools['workflow_communication'])
        
        return filled_df
    
    def generate_documents(self, df, output_folder="output_documents_completo"):
        """Gera documentos Word para cada shift"""
        Path(output_folder).mkdir(exist_ok=True)
        successful_docs = 0
        
        for index, row in df.iterrows():
            try:
                # Preparar dados do workflow
                workflow_data = []
                for i in range(1, 9):
                    timeline = row.get(f'workflow_timeline_{i}', '')
                    task = row.get(f'workflow_task_{i}', '')
                    equipment = row.get(f'workflow_equipment_{i}', '')
                    communication = row.get(f'workflow_communication_{i}', '')
                    
                    if timeline and task:  # Só adiciona se tiver dados
                        workflow_data.append({
                            'timeline': timeline,
                            'task': task,
                            'equipment': equipment,
                            'communication': communication
                        })
                
                # Contexto para o template
                context = {
                    'student_name': self.config.get('student_name', 'Student Name'),
                    'supervisor': self.config.get('supervisor', 'Supervisor Name'),
                    'course': self.config.get('course', 'SIT40521 Certificate IV in Kitchen Management'),
                    'shift_number': row['shift_number'],
                    'establishment': row['establishment'],
                    'date': row['date'],
                    'shift_type': row['shift_type'].title(),
                    'start_time': row['start_time'],
                    'end_time': row['end_time'],
                    'hours': row['hours'],
                    'menu_style': row['menu_style'],
                    'prepared_for_service': row['prepared_for_service'],
                    'special_requests': row['special_requests'],
                    'food_details': row['food_details'],
                    'complaints_problems': row['complaints_problems'],
                    'solutions_implemented': row['solutions_implemented'],
                    'debrief_learnings': row['debrief_learnings'],
                    'handover_completed': row['handover_completed'],
                    'customer_feedback': row['customer_feedback'],
                    'workflow_table': workflow_data
                }
                
                # Nome do arquivo de saída
                output_filename = f"shift_{row['shift_number']:02d}_{row['date']}_{row['shift_type']}.docx"
                output_path = os.path.join(output_folder, output_filename)
                
                # Gerar documento
                self.create_document(context, output_path)
                successful_docs += 1
                
            except Exception as e:
                print(f"❌ Erro ao gerar documento para shift {row['shift_number']}: {e}")
        
        print(f"✅ {successful_docs} documentos gerados com sucesso em '{output_folder}'")
    
    def create_document(self, context, output_path):
        """Cria documento Word individual"""
        from docx import Document
        
        doc = Document()
        
        # Título
        title = doc.add_heading(f'Logbook Entry - Shift {context["shift_number"]}', 0)
        
        # Informações básicas
        doc.add_paragraph(f"Student: {context['student_name']}")
        doc.add_paragraph(f"Course: {context['course']}")
        doc.add_paragraph(f"Supervisor: {context['supervisor']}")
        doc.add_paragraph(f"Name of establishment: {context['establishment']}")
        doc.add_paragraph(f"Date of shift: {context['date']}")
        doc.add_paragraph(f"Start & finish time of shift: {context['start_time']} - {context['end_time']}")
        doc.add_paragraph(f"Hours: {context['hours']}")
        doc.add_paragraph(f"Shift type: ☑ {context['shift_type']}")
        doc.add_paragraph(f"Menu/food service style: ☑ {context['menu_style']}")
        
        # Seções principais
        doc.add_heading('Prepared for service', level=1)
        doc.add_paragraph(context['prepared_for_service'])
        
        doc.add_heading('Special requests', level=1)
        doc.add_paragraph(context['special_requests'])
        
        doc.add_heading('Food details', level=1)
        doc.add_paragraph(context['food_details'])
        
        doc.add_heading('Customer issues/complaints', level=1)
        doc.add_paragraph(context['complaints_problems'])
        
        doc.add_heading('Solutions implemented', level=1)
        doc.add_paragraph(context['solutions_implemented'])
        
        doc.add_heading('Debrief summary', level=1)
        doc.add_paragraph(context['debrief_learnings'])
        
        doc.add_heading('Successful handover completed', level=1)
        doc.add_paragraph(context['handover_completed'])
        
        doc.add_heading('Customer feedback', level=1)
        doc.add_paragraph(context['customer_feedback'])
        
        # Workflow table
        doc.add_heading(f'Workflow plan - {context["shift_type"]} Shift', level=1)
        
        # Criar tabela do workflow
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Cabeçalhos
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Timeline'
        hdr_cells[1].text = 'Task (Description) and Priority'
        hdr_cells[2].text = 'Equipment & WHS'
        hdr_cells[3].text = 'Communication (Who, About what?)'
        
        # Adicionar dados do workflow
        for workflow_item in context['workflow_table']:
            row_cells = table.add_row().cells
            row_cells[0].text = workflow_item['timeline']
            row_cells[1].text = workflow_item['task']
            row_cells[2].text = workflow_item['equipment']
            row_cells[3].text = workflow_item['communication']
        
        doc.save(output_path)

def main():
    """Função principal do script"""
    parser = argparse.ArgumentParser(description='Processar CSV completo e gerar documentos')
    parser.add_argument('--csv', default='PLANILHA_COMPLETA_LOGBOOK.csv', help='Caminho para arquivo CSV preenchido')
    parser.add_argument('--output', default='output_documents_completo', help='Pasta de saída para documentos')
    parser.add_argument('--config', default='config.json', help='Arquivo de configuração')
    parser.add_argument('--pools', default='content_pools.json', help='Arquivo de pools de conteúdo')
    
    args = parser.parse_args()
    
    print("🚀 Iniciando processamento de CSV completo...")
    
    # Inicializa processador
    processor = CSVProcessorCompleto(args.config, args.pools)
    
    # Carrega dados CSV
    df = processor.load_csv_data(args.csv)
    if df is None:
        return
    
    # Preenche campos vazios
    print("🔄 Preenchendo campos vazios automaticamente...")
    df_filled = processor.fill_missing_data(df)
    
    # Gera documentos
    print("📄 Gerando documentos Word completos...")
    processor.generate_documents(df_filled, args.output)
    
    print("✅ Processamento concluído com sucesso!")
    print(f"📁 Documentos salvos em: {args.output}")

if __name__ == "__main__":
    main() 
